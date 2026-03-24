#!/usr/bin/env python3
"""
검토보고서 DOCX 생성기
JSON 데이터를 받아 맑은고딕 기반의 Word 문서를 생성합니다.

사용법:
    python generate_docx.py <input_json_path> <output_docx_path>
"""

import json
import sys
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


FONT_NAME = "맑은 고딕"


def set_cell_shading(cell, color_hex):
    """셀 배경색 설정"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, size=10, bold=False, color=None):
    """런(텍스트) 폰트 설정"""
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading_styled(doc, text, level=1):
    """맑은고딕 스타일 적용된 헤딩 추가"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    size_map = {1: 18, 2: 14, 3: 12}
    run.font.size = Pt(size_map.get(level, 10))
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return heading


def add_paragraph_styled(doc, text, size=10, bold=False, color=None):
    """맑은고딕 스타일 적용된 문단 추가"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size, bold, color)
    return p


def create_table(doc, headers, rows, col_widths=None):
    """스타일 적용된 테이블 생성"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    header_row = table.rows[0]
    for i, header_text in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header_text)
        set_run_font(run, size=9, bold=True, color=(0xFF, 0xFF, 0xFF))
        set_cell_shading(cell, "2D3748")

    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            set_run_font(run, size=9)
            if row_idx % 2 == 0:
                set_cell_shading(cell, "F7FAFC")

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


def get_status_color(score):
    """점수에 따른 색상 반환"""
    if score >= 10.0:
        return (0x27, 0xAE, 0x60)
    elif score >= 7.0:
        return (0x29, 0x80, 0xB9)
    elif score >= 4.0:
        return (0xF3, 0x9C, 0x12)
    elif score >= 1.0:
        return (0xE7, 0x4C, 0x3C)
    else:
        return (0x95, 0xA5, 0xA6)


def generate_docx(data, output_path):
    """DOCX 문서 생성 메인 함수"""
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(10)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    report = data.get("report", {})
    items = data.get("items", [])
    phase1_history = data.get("phase1_history", [])

    # 제목
    title = doc.add_heading(level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("지시서-작업물 대조 검토 보고서")
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"생성일시: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    set_run_font(run, size=9, color=(0x71, 0x71, 0x71))

    doc.add_paragraph()

    # 1. 개요
    add_heading_styled(doc, "1. 개요", level=1)

    overview_rows = [
        ["지시서 파일", report.get("instruction_file", "-")],
        ["작업물 파일", report.get("deliverable_file", "-")],
        ["Phase 1 분석 신뢰도", f"{report.get('phase1_score', '-')} / 10.0"],
        ["Phase 2 종합 점수", f"{report.get('phase2_overall_score', '-')} / 10.0"],
    ]
    create_table(doc, ["항목", "내용"], overview_rows, col_widths=[5, 12])
    doc.add_paragraph()

    add_heading_styled(doc, "요약", level=3)
    summary_rows = [
        ["충족 (10.0)", str(report.get("fulfilled_count", 0))],
        ["대부분충족 (7.0~9.9)", str(report.get("partial_high_count", 0))],
        ["부분충족 (4.0~6.9)", str(report.get("partial_count", 0))],
        ["미흡 (1.0~3.9)", str(report.get("insufficient_count", 0))],
        ["누락 (0.0)", str(report.get("missing_count", 0))],
    ]
    create_table(doc, ["상태", "항목 수"], summary_rows, col_widths=[6, 4])
    doc.add_paragraph()

    # 2. 항목별 대조 결과
    add_heading_styled(doc, "2. 항목별 대조 결과", level=1)

    detail_headers = ["#", "분류", "요구항목", "1차", "2차", "최종", "상태", "태그"]
    detail_rows = []
    for item in items:
        tag = item.get("tag") or "-"
        detail_rows.append([
            str(item.get("id", "")),
            item.get("category", ""),
            item.get("requirement", ""),
            str(item.get("score_r1", "")),
            str(item.get("score_r2", "")),
            str(item.get("final_score", "")),
            item.get("status", ""),
            tag,
        ])

    if detail_rows:
        create_table(
            doc, detail_headers, detail_rows,
            col_widths=[0.8, 2.5, 4.5, 1.2, 1.2, 1.2, 2, 1.2]
        )
    doc.add_paragraph()

    # 3. 보완 필요 체크리스트
    add_heading_styled(doc, "3. 보완 필요 체크리스트", level=1)
    add_paragraph_styled(doc, "아래는 최종 점수 9.9점 이하인 항목들입니다.", size=10)
    doc.add_paragraph()

    score_groups = [
        ("3-1. 누락 항목 (0.0점)", lambda x: x.get("final_score", 0) == 0.0),
        ("3-2. 미흡 항목 (1.0~3.9점)", lambda x: 1.0 <= x.get("final_score", 0) <= 3.9),
        ("3-3. 부분충족 항목 (4.0~6.9점)", lambda x: 4.0 <= x.get("final_score", 0) <= 6.9),
        ("3-4. 대부분충족 항목 (7.0~9.9점)", lambda x: 7.0 <= x.get("final_score", 0) <= 9.9),
    ]

    for group_title, filter_fn in score_groups:
        group_items = [item for item in items if filter_fn(item)]
        if not group_items:
            continue

        add_heading_styled(doc, group_title, level=2)

        for item in group_items:
            tag_str = f" [{item.get('tag')}]" if item.get("tag") else ""
            score = item.get("final_score", 0)

            p = doc.add_paragraph()
            run = p.add_run(f"  {item.get('requirement', '')}")
            set_run_font(run, size=11, bold=True)
            run = p.add_run(f"  ({score}/10.0){tag_str}")
            set_run_font(run, size=10, bold=False, color=get_status_color(score))

            if item.get("current_state"):
                p = doc.add_paragraph()
                run = p.add_run("    현재 상태: ")
                set_run_font(run, size=9, bold=True, color=(0x55, 0x55, 0x55))
                run = p.add_run(item["current_state"])
                set_run_font(run, size=9)

            if item.get("suggestion_direction"):
                p = doc.add_paragraph()
                run = p.add_run("    보완 방향: ")
                set_run_font(run, size=9, bold=True, color=(0x29, 0x80, 0xB9))
                run = p.add_run(item["suggestion_direction"])
                set_run_font(run, size=9)

            if item.get("suggestion_example"):
                p = doc.add_paragraph()
                run = p.add_run("    작성 예시: ")
                set_run_font(run, size=9, bold=True, color=(0x27, 0xAE, 0x60))
                run = p.add_run(f'"{item["suggestion_example"]}"')
                set_run_font(run, size=9)

            doc.add_paragraph()

    # 4. 충족 항목
    fulfilled_items = [item for item in items if item.get("final_score", 0) >= 10.0]
    if fulfilled_items:
        add_heading_styled(doc, "4. 충족 항목", level=1)
        fulfilled_rows = [
            [str(item.get("id", "")), item.get("category", ""),
             item.get("requirement", ""), str(item.get("final_score", ""))]
            for item in fulfilled_items
        ]
        create_table(
            doc, ["#", "분류", "요구항목", "최종점수"], fulfilled_rows,
            col_widths=[1, 3, 8, 2]
        )
    doc.add_paragraph()

    # 5. Phase 1 분석 이력
    add_heading_styled(doc, "5. Phase 1 분석 이력", level=1)
    if phase1_history:
        history_rows = [
            [h.get("round", ""), str(h.get("score", ""))]
            for h in phase1_history
        ]
        create_table(doc, ["단계", "점수"], history_rows, col_widths=[6, 4])

    doc.save(output_path)
    print(f"DOCX 생성 완료: {output_path}")


def main():
    if len(sys.argv) != 3:
        print("사용법: python generate_docx.py <input_json> <output_docx>")
        sys.exit(1)

    json_path = sys.argv[1]
    output_path = sys.argv[2]

    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    generate_docx(data, output_path)


if __name__ == "__main__":
    main()
