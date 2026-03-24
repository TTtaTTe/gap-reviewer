#!/usr/bin/env python3
"""
수정이력 보고서 및 수정본 DOCX 생성기
gap-fix의 수정 이력 JSON과 수정본 텍스트를 받아 Word 문서를 생성합니다.

사용법:
    python generate_fix_docx.py <modification_json_path> <content_text_path> <output_docx_path>

    modification_json_path: 수정 이력 JSON 파일 경로
    content_text_path: 수정본 텍스트 파일 경로 (없으면 "none" 전달)
    output_docx_path: 출력 DOCX 파일 경로
"""

import json
import sys
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


FONT_NAME = "맑은 고딕"


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, size=10, bold=False, color=None):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    size_map = {1: 18, 2: 14, 3: 12}
    run.font.size = Pt(size_map.get(level, 10))
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return heading


def add_paragraph_styled(doc, text, size=10, bold=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size, bold, color)
    return p


def create_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    header_row = table.rows[0]
    for i, header_text in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header_text)
        set_run_font(run, size=9, bold=True, color=(0xFF, 0xFF, 0xFF))
        set_cell_shading(cell, "2D3748")

    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            set_run_font(run, size=9)
            if row_idx % 2 == 0:
                set_cell_shading(cell, "F7FAFC")

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


def generate_modification_report(data, output_path):
    """수정 이력 보고서 DOCX 생성"""
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(10)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    log = data.get("modification_log", {})
    modifications = log.get("modifications", [])
    remaining = log.get("remaining_items", [])

    # 제목
    title = doc.add_heading(level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("수정 이력 보고서")
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"생성일시: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    set_run_font(run, size=9, color=(0x71, 0x71, 0x71))
    doc.add_paragraph()

    # 1. 개요
    add_heading_styled(doc, "1. 개요", level=1)
    overview_rows = [
        ["검토보고서", log.get("source_json", "-")],
        ["원본 작업물", log.get("original_file", "-")],
        ["수정본 파일", log.get("output_file", "-")],
        ["전체 보완 필요", f"{log.get('total_gaps', 0)}개"],
        ["수정 완료", f"{log.get('fixed_count', 0)}개"],
        ["잔여 항목", f"{log.get('remaining_count', 0)}개"],
    ]
    create_table(doc, ["항목", "내용"], overview_rows, col_widths=[5, 12])
    doc.add_paragraph()

    # 2. 수정 항목 상세
    add_heading_styled(doc, "2. 수정 항목 상세", level=1)
    if modifications:
        for mod in modifications:
            p = doc.add_paragraph()
            run = p.add_run(f"  #{mod.get('item_id', '')} {mod.get('requirement', '')}")
            set_run_font(run, size=11, bold=True)
            run = p.add_run(f"  (기존 {mod.get('before_score', 0)}점)")
            set_run_font(run, size=10, color=(0xE7, 0x4C, 0x3C))

            p = doc.add_paragraph()
            run = p.add_run("    수정 내용: ")
            set_run_font(run, size=9, bold=True, color=(0x29, 0x80, 0xB9))
            run = p.add_run(f"{mod.get('action', '')} — {mod.get('description', '')}")
            set_run_font(run, size=9)

            if mod.get("cover_source"):
                p = doc.add_paragraph()
                run = p.add_run("    출처: ")
                set_run_font(run, size=9, bold=True, color=(0x55, 0x55, 0x55))
                run = p.add_run(mod["cover_source"])
                set_run_font(run, size=9)

            doc.add_paragraph()
    else:
        add_paragraph_styled(doc, "수정된 항목이 없습니다.", size=10)

    # 3. 미수정 잔여 항목
    add_heading_styled(doc, "3. 미수정 잔여 항목", level=1)
    if remaining:
        remaining_rows = [
            [
                str(item.get("item_id", "")),
                item.get("requirement", ""),
                str(item.get("final_score", "")),
                item.get("reason", ""),
            ]
            for item in remaining
        ]
        create_table(
            doc, ["#", "요구항목", "점수", "미수정 사유"],
            remaining_rows, col_widths=[1, 5, 1.5, 8]
        )
    else:
        add_paragraph_styled(doc, "모든 보완 필요 항목이 수정되었습니다.", size=10)

    doc.save(output_path)
    print(f"수정이력 DOCX 생성 완료: {output_path}")


def generate_fixed_document(content_text, output_path):
    """수정본 DOCX 생성 (텍스트 기반)"""
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(10)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    for line in content_text.split("\n"):
        stripped = line.strip()
        if stripped.startswith("# "):
            add_heading_styled(doc, stripped[2:], level=1)
        elif stripped.startswith("## "):
            add_heading_styled(doc, stripped[3:], level=2)
        elif stripped.startswith("### "):
            add_heading_styled(doc, stripped[4:], level=3)
        elif stripped == "":
            doc.add_paragraph()
        else:
            p = doc.add_paragraph()
            if "[수정됨]" in stripped:
                parts = stripped.split("[수정됨]")
                run = p.add_run(parts[0])
                set_run_font(run, size=10)
                run = p.add_run("[수정됨]")
                set_run_font(run, size=9, bold=True, color=(0x27, 0xAE, 0x60))
                if len(parts) > 1:
                    run = p.add_run(parts[1])
                    set_run_font(run, size=10)
            else:
                run = p.add_run(stripped)
                set_run_font(run, size=10)

    doc.save(output_path)
    print(f"수정본 DOCX 생성 완료: {output_path}")


def main():
    if len(sys.argv) != 4:
        print("사용법: python generate_fix_docx.py <modification_json> <content_text> <output_docx>")
        print("  content_text에 'none'을 전달하면 수정이력 보고서만 생성합니다.")
        sys.exit(1)

    json_path = sys.argv[1]
    content_path = sys.argv[2]
    output_path = sys.argv[3]

    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    if content_path.lower() == "none":
        generate_modification_report(data, output_path)
    else:
        with open(content_path, "r", encoding="utf-8") as f:
            content_text = f.read()
        generate_fixed_document(content_text, output_path)


if __name__ == "__main__":
    main()
