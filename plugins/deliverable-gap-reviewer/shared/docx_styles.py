#!/usr/bin/env python3
"""
공통 DOCX 스타일 모듈
검토보고서(gap-review)와 수정본/수정이력(gap-fix) 모두에서 사용.
"""

import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement


# ── 색상 팔레트 ──
COLOR_NAVY = RGBColor(0x1B, 0x2A, 0x4A)
COLOR_CHARCOAL = RGBColor(0x33, 0x33, 0x33)
COLOR_BODY = RGBColor(0x40, 0x40, 0x40)
COLOR_GRAY = RGBColor(0x80, 0x80, 0x80)
COLOR_LIGHT_GRAY = RGBColor(0x99, 0x99, 0x99)
COLOR_SUBTITLE = RGBColor(0x66, 0x66, 0x66)
COLOR_GREEN = RGBColor(0x27, 0xAE, 0x60)
COLOR_RED = RGBColor(0xE7, 0x4C, 0x3C)
COLOR_BLUE = RGBColor(0x29, 0x80, 0xB9)
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

HEX_TABLE_HEADER = "2C3E50"
HEX_TABLE_ALT_ROW = "F2F2F2"
HEX_DIVIDER = "D9D9D9"

FONT_NAME = "맑은 고딕"


# ── 기본 문서 생성 ──

def create_document():
    """기본 설정이 적용된 Document 생성"""
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(11)
    font.color.rgb = COLOR_BODY
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.15

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    return doc


# ── 폰트 설정 ──

def set_run_font(run, size=11, bold=False, color=None, italic=False):
    """런 폰트 설정"""
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


# ── 표지 페이지 ──

def add_cover_page(doc, title, subtitle=None, date_str=None,
                   author=None, classification=None, version=None):
    """전문 표지 페이지 생성"""
    # 상단 여백
    for _ in range(6):
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(0)
        spacer.paragraph_format.space_before = Pt(0)

    # 제목
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(title)
    set_run_font(run, size=24, bold=True, color=COLOR_NAVY)

    # 부제
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(20)
        run = p.add_run(subtitle)
        set_run_font(run, size=14, color=COLOR_SUBTITLE)

    # 구분선
    _add_horizontal_rule(doc)

    # 날짜
    if date_str:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(date_str)
        set_run_font(run, size=11, color=COLOR_GRAY)

    # 작성자
    if author:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(author)
        set_run_font(run, size=11, color=COLOR_GRAY)

    # 하단 여백
    for _ in range(4):
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(0)
        spacer.paragraph_format.space_before = Pt(0)

    # 문서분류 / 버전
    meta_parts = []
    if classification:
        meta_parts.append(f"문서분류: {classification}")
    if version:
        meta_parts.append(f"버전: {version}")
    if meta_parts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(" │ ".join(meta_parts))
        set_run_font(run, size=10, color=COLOR_LIGHT_GRAY)

    # 페이지 나누기
    doc.add_page_break()


def _add_horizontal_rule(doc):
    """구분선 추가"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), HEX_DIVIDER)
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── 목차 ──

def add_toc(doc):
    """자동 목차 필드 삽입 (Word에서 F9로 업데이트 필요)"""
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = h.add_run("목차")
    set_run_font(run, size=16, bold=True, color=COLOR_NAVY)

    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(4)
    run2 = p2.add_run("(Word에서 이 영역을 우클릭 → '필드 업데이트'로 목차를 생성하세요)")
    set_run_font(run2, size=9, italic=True, color=COLOR_GRAY)

    doc.add_page_break()


# ── 헤딩 ──

def add_heading(doc, text, level=1):
    """스타일 적용된 헤딩"""
    heading = doc.add_heading(level=level)
    heading.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    heading.paragraph_format.space_after = Pt(8)
    run = heading.add_run(text)
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)

    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = COLOR_NAVY
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = COLOR_CHARCOAL
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_CHARCOAL

    run.font.bold = True
    return heading


# ── 본문 ──

def add_paragraph(doc, text, size=11, bold=False, color=None, alignment=None):
    """스타일 적용된 문단"""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment

    # [수정됨] 태그 처리
    if "[수정됨]" in text:
        parts = text.split("[수정됨]")
        run = p.add_run(parts[0])
        set_run_font(run, size=size, bold=bold, color=color or COLOR_BODY)
        run = p.add_run("[수정됨]")
        set_run_font(run, size=max(size - 1, 9), bold=True, color=COLOR_GREEN)
        if len(parts) > 1 and parts[1]:
            run = p.add_run(parts[1])
            set_run_font(run, size=size, bold=bold, color=color or COLOR_BODY)
    else:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color or COLOR_BODY)

    return p


# ── 체크리스트 항목 ──

def add_checklist_item(doc, title, score, tag=None, current_state=None,
                       suggestion_direction=None, suggestion_example=None):
    """보완 필요 체크리스트 항목"""
    score_color = _get_score_color(score)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("□  ")
    set_run_font(run, size=11, color=COLOR_CHARCOAL)
    run = p.add_run(title)
    set_run_font(run, size=11, bold=True, color=COLOR_CHARCOAL)
    run = p.add_run(f"  ({score}/10.0)")
    set_run_font(run, size=10, color=score_color)
    if tag:
        run = p.add_run(f"  [{tag}]")
        set_run_font(run, size=9, bold=True, color=COLOR_BLUE)

    if current_state:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run("현재 상태: ")
        set_run_font(run, size=10, bold=True, color=COLOR_GRAY)
        run = p.add_run(current_state)
        set_run_font(run, size=10, color=COLOR_BODY)

    if suggestion_direction:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run("보완 방향: ")
        set_run_font(run, size=10, bold=True, color=COLOR_BLUE)
        run = p.add_run(suggestion_direction)
        set_run_font(run, size=10, color=COLOR_BODY)

    if suggestion_example:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run("작성 예시: ")
        set_run_font(run, size=10, bold=True, color=COLOR_GREEN)
        run = p.add_run(f'"{suggestion_example}"')
        set_run_font(run, size=10, color=COLOR_BODY)


def _get_score_color(score):
    if score >= 10.0:
        return COLOR_GREEN
    elif score >= 7.0:
        return COLOR_BLUE
    elif score >= 4.0:
        return RGBColor(0xF3, 0x9C, 0x12)
    elif score >= 1.0:
        return COLOR_RED
    else:
        return COLOR_GRAY


# ── 표 ──

def set_cell_shading(cell, color_hex):
    """셀 배경색 설정"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def create_table(doc, headers, rows, col_widths=None):
    """전문 스타일 테이블 생성"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # 헤더 행
    header_row = table.rows[0]
    for i, header_text in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header_text)
        set_run_font(run, size=10, bold=True, color=COLOR_WHITE)
        set_cell_shading(cell, HEX_TABLE_HEADER)

    # 데이터 행
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            set_run_font(run, size=10, color=COLOR_CHARCOAL)
            if row_idx % 2 == 1:
                set_cell_shading(cell, HEX_TABLE_ALT_ROW)

    # 열 너비
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()  # 표 후 간격
    return table


# ── 마크다운 표 파서 ──

def parse_md_table(lines):
    """마크다운 표 라인들을 (headers, rows) 튜플로 파싱"""
    if len(lines) < 2:
        return None, None

    def split_row(line):
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        return cells

    headers = split_row(lines[0])

    # 구분선(---|---) 건너뛰기
    start_idx = 1
    if start_idx < len(lines) and re.match(r"^\|?[\s\-:|]+\|", lines[start_idx]):
        start_idx = 2

    rows = []
    for line in lines[start_idx:]:
        if "|" in line:
            rows.append(split_row(line))

    return headers, rows


# ── 머리글/꼬리글 ──

def add_header_footer(doc, header_text="", classification=""):
    """머리글(문서제목) + 꼬리글(페이지번호) 설정"""
    for section in doc.sections:
        section.different_first_page_header_footer = True

        # 머리글 (첫 페이지 제외)
        header = section.header
        header.is_linked_to_previous = False
        if header.paragraphs:
            hp = header.paragraphs[0]
        else:
            hp = header.add_paragraph()
        hp.clear()

        if header_text:
            run = hp.add_run(header_text)
            set_run_font(run, size=9, color=COLOR_GRAY)
        if classification:
            run = hp.add_run(f"\t\t{classification}")
            set_run_font(run, size=9, color=COLOR_GRAY)

        # 머리글 하단 테두리
        pPr = hp._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), HEX_DIVIDER)
        pBdr.append(bottom)
        pPr.append(pBdr)

        # 꼬리글 (페이지 번호)
        footer = section.footer
        footer.is_linked_to_previous = False
        if footer.paragraphs:
            fp = footer.paragraphs[0]
        else:
            fp = footer.add_paragraph()
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_page_number(fp)


def _add_page_number(paragraph):
    """페이지 번호 필드 삽입"""
    run = paragraph.add_run()
    set_run_font(run, size=9, color=COLOR_GRAY)

    run._r.append(_make_fld_char("begin"))

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)

    run._r.append(_make_fld_char("separate"))
    run._r.append(_make_fld_char("end"))


def _make_fld_char(fld_type):
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), fld_type)
    return fld


# ── 마크다운 → DOCX 변환 ──

def md_to_docx(doc, md_text):
    """마크다운 텍스트를 DOCX 요소로 변환"""
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 빈 줄
        if not stripped:
            i += 1
            continue

        # 마크다운 표 감지
        if "|" in stripped and i + 1 < len(lines):
            table_lines = []
            j = i
            while j < len(lines) and "|" in lines[j].strip():
                table_lines.append(lines[j])
                j += 1
            if len(table_lines) >= 2:
                headers, rows = parse_md_table(table_lines)
                if headers and rows:
                    create_table(doc, headers, rows)
                    i = j
                    continue

        # 헤딩
        if stripped.startswith("### "):
            add_heading(doc, stripped[4:], level=3)
        elif stripped.startswith("## "):
            add_heading(doc, stripped[3:], level=2)
        elif stripped.startswith("# "):
            add_heading(doc, stripped[2:], level=1)
        # 수평선
        elif stripped in ("---", "***", "___"):
            _add_horizontal_rule(doc)
        # 체크박스
        elif stripped.startswith("- [ ] ") or stripped.startswith("- [x] "):
            add_paragraph(doc, stripped, size=10)
        # 볼드 라인
        elif stripped.startswith("**") and stripped.endswith("**"):
            add_paragraph(doc, stripped[2:-2], bold=True)
        # 인용문
        elif stripped.startswith("> "):
            p = add_paragraph(doc, stripped[2:], size=10, color=COLOR_GRAY)
            p.paragraph_format.left_indent = Cm(1)
        # 일반 본문
        else:
            add_paragraph(doc, stripped)

        i += 1
